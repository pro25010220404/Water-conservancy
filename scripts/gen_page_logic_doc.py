# -*- coding: utf-8 -*-
"""三页面：按钮逻辑 + 规则 + 接口需求 — 给后端"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

DESKTOP = os.path.join(os.path.expanduser('~'), 'Desktop')
OUTPUT = os.path.join(DESKTOP, '三页面-按钮逻辑与接口说明（给后端）.docx')


def font(run, size=10.5, bold=False):
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def cell(c, text, bold=False):
    c.text = ''
    font(c.paragraphs[0].add_run(str(text)), bold=bold)


def tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell(t.rows[0].cells[i], h, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell(t.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = '黑体'
        r.font.size = Pt(13 if level == 1 else 11)
        r.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def p(doc, text):
    para = doc.add_paragraph()
    font(para.add_run(text))


def bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    font(para.add_run(text))


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = s.right_margin = Cm(2.8)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('三页面 · 按钮逻辑 / 规则 / 接口说明')
    r.font.name = '黑体'
    r.font.size = Pt(17)
    r.bold = True
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(sub.add_run(f'给后端 · {datetime.now():%Y-%m-%d}'), 10)

    p(doc, '每页结构：① 按钮怎么做  ② 业务规则  ③ 要调哪些接口（✓已有 / ★待开发）')

    # ═══════════════════════════════════════
    h(doc, '0. 全局')
    h(doc, '0.1 三页职责', 2)
    tbl(doc, ['页面', '路由', '干什么'], [
        ['虚拟仿真', '/virtual-simulation', '手调水位/降雨 → 预览 KPI → 应用后全站显示仿真值'],
        ['节点控制', '/dispatch/gates', '12 孔设 target → 互锁 → 执行 → opening 更新'],
        ['数字孪生', '/simulation', '3D 仿真任务 + AI 模型 + 报告（独立链路）'],
    ])

    h(doc, '0.2 全局规则', 2)
    tbl(doc, ['规则', '说明'], [
        ['手动模式', 'mode=manual 或 autoLevel=1 才能调闸；否则节点控制全禁用'],
        ['opening vs target', 'opening=设备当前开度；target=用户待下发；两者不等=待提交'],
        ['离线孔', 'status=offline 不可调不可执行（11# 12# 底孔默认离线）'],
        ['开度范围', '0–100%，步长 1%'],
        ['轮询', '节点控制每 5s 拉 gates+realtime；有未提交 target 时前端保留本地 target'],
        ['响应格式', '{ code:0, msg, data, trace_id }；鉴权 Bearer token；默认 reservoir_id=1'],
        ['演示绕过互锁', '仅前端 sessionStorage；生产后端必须始终校验互锁，不提供绕过'],
    ])

    h(doc, '0.3 互锁规则（节点 execute 必校验）', 2)
    tbl(doc, ['rule_code', '规则', '阈值', '级别', '后端行为'], [
        ['symmetry_constraint', '表孔对称', '1#–8# 在线孔 target 最大差 ≤10%', 'block', '拒绝执行，返回 violations'],
        ['rate_exceeded', '单步变化', '|target−opening| ≤10%', 'warn', '可执行，返回 warnings'],
        ['min_discharge_guarantee', '生态流量', '预估出库 ≥500 m³/s', 'warn', '可执行，返回 warnings'],
    ])
    p(doc, 'violations 建议格式：{ passed, violations:[{ rule_code, rule_name, severity, affected_gate_ids, message }] }')

    # ═══════════════════════════════════════
    h(doc, '1. 虚拟仿真页')
    p(doc, '布局：顶快捷按钮 | 左参数区 | 右预览（水位柱+流量+坝体12孔）')

    h(doc, '1.1 按钮逻辑 + 规则 + 接口', 2)
    tbl(doc, ['按钮/操作', '规则（前置）', '点击后逻辑', '接口'], [
        ['常水位/汛限偏高/枯水偏低', 'locked=false', '设预设水位+降雨 → 重算预览', '无（纯前端）'],
        ['上游水位滑块', 'locked=false；110–200m', '改水位 → 重算', '无'],
        ['降雨量滑块', 'locked=false；0–120mm/h', '改降雨 → 重算', '无'],
        ['重置参数', '无', 'active=false；恢复 baseline', '✓ GET realtime（取基准）'],
        ['锁定模拟', '无', 'locked=true，禁用滑块', '无'],
        ['解锁调节', 'locked=true', 'locked=false', '无'],
        ['一键应用仿真', '无', 'active=true；全站 overlay 仿真 KPI', '无（前端标记）'],
        ['节点控制', '无', '跳 /dispatch/gates', '—'],
        ['数字孪生', '无', '跳 /simulation', '—'],
        ['去节点控制调闸', 'manual 或 L1', '否则跳运行控制', '✓ 读 mode（现 Mock）'],
        ['坝体点孔', '无', '选中 gateId', '✓ GET gates'],
    ])

    h(doc, '1.2 本页业务规则', 2)
    bullet(doc, '拖滑块 = 仅预览；必须「一键应用仿真」其他页才变')
    bullet(doc, '计算在前端 virtualSimulationEngine，与数字孪生仿真任务无关')
    bullet(doc, 'active 后 fetchRealtimeKpi / fetchGates 返回值被 overlay，不改真实设备')
    bullet(doc, '进页调 refreshCore → 读 realtime+gates 初始化 baseline')

    h(doc, '1.3 接口需求', 2)
    tbl(doc, ['状态', '方法', '路径', '说明'], [
        ['✓ 已有', 'GET', '/api/monitoring/realtime?reservoir_id=1', '基准水位/流量'],
        ['✓ 已有', 'GET', '/api/monitoring/gates?reservoir_id=1', '12 孔开度'],
        ['★ 可选', 'POST', '/api/v1/hydrology/virtual-sim/apply', 'body:{upstream_level,rainfall}→衍生KPI'],
        ['★ 可选', 'DELETE', '/api/v1/hydrology/virtual-sim', '重置仿真工况'],
    ])

    # ═══════════════════════════════════════
    h(doc, '2. 节点控制页')
    p(doc, '布局：顶操作栏 | 坝体剖面 | 左闸门列表 | 右孔详情 | 底待提交表')

    h(doc, '2.1 按钮逻辑 + 规则 + 接口', 2)
    tbl(doc, ['按钮', '规则（前置）', '点击后逻辑', '接口'], [
        ['互锁严格/手动绕过', '演示功能', '切换能否强制提交', '无（生产后端不提供）'],
        ['虚拟仿真', '无', '跳转', '—'],
        ['数字孪生', '无', '跳转', '—'],
        ['表孔拉齐', 'manual；有在线表孔', '1#–8# target 设为同值', '无（前端算）'],
        ['全部复位', 'manual', '全部 target:=opening', '无'],
        ['提交变更(N)', 'manual；无 block；N>0', '批量执行 N 孔', '★ POST .../gates/batch-execute'],
        ['去运行控制', '非 manual 时显示', '跳转切模式', '★ PUT .../dispatch/mode'],
        ['详情', '有 violations', '弹窗', '★ POST .../gates/precheck（可选）'],
        ['调参', '仿真 active', '跳虚拟仿真', '—'],
        ['选孔/滑块/上下调', 'manual+在线', '改 target；重算互锁', '✓ GET gates（轮询）'],
        ['执行本孔', 'manual+在线+互锁过', '单孔 opening:=target', '★ POST .../gates/{id}/execute'],
        ['重置', 'manual', '该孔 target:=opening', '无'],
        ['同组同步', 'manual+在线', '同组 target 对齐', '无'],
    ])

    h(doc, '2.2 本页业务规则', 2)
    bullet(doc, 'canManualControl = mode=manual OR autoLevel=1')
    bullet(doc, 'pendingChanges = 在线且 target≠opening 的孔数')
    bullet(doc, '执行流程：确认 →（互锁 bypass 确认）→ API 下发 → status:executing → opening:=target → online')
    bullet(doc, '离开 Tab 且 pendingChanges>0 → 弹窗确认')
    bullet(doc, 'gate 分组：001–008 表孔 / 009–010 中孔 / 011–012 底孔')

    h(doc, '2.3 gates 单条字段（GET /monitoring/gates）', 2)
    tbl(doc, ['字段', '类型', '说明'], [
        ['id', 'number', '设备 id'],
        ['code', 'string', 'GATE-001 ~ GATE-012'],
        ['name', 'string', '1# ~ 底2#'],
        ['status', 'string', 'online | offline | executing'],
        ['opening', 'number', '当前开度 0–100'],
        ['target_opening', 'number', '目标开度；应与 opening 分开维护'],
        ['mode', 'string', 'AI-DQN / manual 等'],
        ['flow_rate', 'number', '单孔流量 m³/s'],
        ['last_action_at', 'string', '最近动作时间'],
    ])

    h(doc, '2.4 接口需求', 2)
    tbl(doc, ['状态', '方法', '路径', '触发场景'], [
        ['✓ 已有', 'GET', '/api/monitoring/gates?reservoir_id=1', '每 5s 轮询'],
        ['✓ 已有', 'GET', '/api/monitoring/realtime?reservoir_id=1', '每 5s 轮询'],
        ['✓ 已有', 'GET', '/api/v1/dispatch/decisions', '进页拼 status'],
        ['✓ 已有', 'POST', '/api/v1/dispatch/execute', '运行控制「总开度执行」'],
        ['★ P0', 'POST', '/api/v1/dispatch/gates/{equipment_id}/execute', '「执行本孔」body:{target_opening}'],
        ['★ P0', 'POST', '/api/v1/dispatch/gates/batch-execute', '「提交变更」body:{items:[{equipment_id,target_opening}]}'],
        ['★ P1', 'POST', '/api/v1/dispatch/gates/precheck', '改 target 后预检，返回 violations'],
        ['★ P1', 'PUT', '/api/v1/dispatch/mode', 'body:{mode:manual|auto}'],
        ['★ P1', 'PUT', '/api/v1/dispatch/auto-level', 'body:{level:1|2|3}'],
        ['★ P2', 'GET', '/api/v1/dispatch/gate-actions', '动作历史'],
        ['★ P2', 'GET', '/api/v1/settings/gate-interlock/rules', '互锁配置（execute 读取）'],
    ])
    p(doc, 'execute 成功后期望：opening=target_opening；写 gate-actions 日志；返回 command_id。')

    # ═══════════════════════════════════════
    h(doc, '3. 数字孪生页')
    p(doc, '布局：左场景库 | 中 2D/3D 视图 | 右 Tab 面板。按钮不跳路由。')

    h(doc, '3.1 按钮逻辑 + 规则 + 接口', 2)
    tbl(doc, ['按钮', '规则（前置）', '点击后逻辑', '接口'], [
        ['2D/3D/BIM', '无', '切换视图/开全景', '无'],
        ['开始仿真', 'idle 或 finished', '启动任务', '✓ POST /v1/simulation/start'],
        ['暂停', 'running', '暂停', '✓ POST /v1/simulation/{id}/pause'],
        ['继续', 'paused', '恢复', '✓ POST /v1/simulation/{id}/resume'],
        ['重置', 'running/paused', '重置', '✓ POST /v1/simulation/{id}/reset'],
        ['调闸门开度', 'running', '400ms 防抖提交', '✓ PUT /v1/simulation/{id}/gate'],
        ['新建场景', '无', '创建', '✓ POST /v1/simulation/scenarios'],
        ['刷新', '无', '拉列表', '✓ GET /v1/simulation/scenarios'],
        ['编辑/删除场景', '删除时 usage_count=0', 'CRUD', '✓ PUT/DELETE .../scenarios/{id}'],
        ['导入模型', '无', '上传', '✓ POST /settings/models/upload'],
        ['激活/训练', '无', '模型操作', '✓ PUT activate + 训练接口'],
        ['生成报告', '有 completed 仿真', '生成 PDF', '✓ POST .../{id}/report'],
        ['下载 PDF', '报告已生成', '下载', '✓ GET 报告下载'],
        ['导入仿真(复盘)', '选中故障', '导入初始条件', '✓ POST import-incident'],
        ['点 1#–5# 孔', '无', '3D 高亮', '无（页内）'],
    ])

    h(doc, '3.2 本页业务规则', 2)
    bullet(doc, '仿真任务与节点控制 execute 是两条独立链路')
    bullet(doc, '不走节点互锁 precheck')
    bullet(doc, 'start 返回 simulation_id + ws_endpoint；前端 WS 推进度，断线 1s 轮询 status')
    bullet(doc, '虚拟仿真 active 时 KPI 显示 overlay 值')
    bullet(doc, '有 usage_count 的场景不可删')

    h(doc, '3.3 start 请求体', 2)
    tbl(doc, ['字段', '说明'], [
        ['scenario_id', '场景 id，默认 1'],
        ['model_id', '模型 id，默认 2'],
        ['reservoir_id', '默认 1'],
        ['duration', '秒，= durationMin×60'],
        ['speed', '1/2/5/10'],
        ['params.initial_water_level', '初始水位 m'],
        ['params.inflow_rate', '入库 m³/s'],
        ['params.gate_opening', '初始开度 0–100'],
    ])

    h(doc, '3.4 接口需求汇总', 2)
    tbl(doc, ['方法', '路径', '说明'], [
        ['GET', '/api/v1/simulation/scenarios', '场景列表'],
        ['POST', '/api/v1/simulation/scenarios', '新建'],
        ['PUT', '/api/v1/simulation/scenarios/{id}', '编辑'],
        ['DELETE', '/api/v1/simulation/scenarios/{id}', '删除'],
        ['POST', '/api/v1/simulation/start', '启动'],
        ['POST', '/api/v1/simulation/{id}/pause|resume|reset', '控制'],
        ['PUT', '/api/v1/simulation/{id}/gate', '{gate_opening}'],
        ['GET', '/api/v1/simulation/{id}/result', '结果'],
        ['POST', '/api/v1/simulation/{id}/report', '报告'],
        ['GET', '/api/v1/simulation/incidents', '故障复盘列表'],
        ['POST', '/api/v1/simulation/import-incident', '{incident_id}'],
        ['GET', '/api/monitoring/realtime', 'KPI 展示'],
    ])

    # ═══════════════════════════════════════
    h(doc, '4. 运行控制页（节点控制前置）')
    tbl(doc, ['按钮', '规则', '逻辑', '接口'], [
        ['自动/手动', '无', '切 mode', '★ PUT /dispatch/mode（现 Mock）'],
        ['总开度执行', 'manual；可提交', '聚合开度一次下发', '✓ POST /dispatch/execute'],
        ['分配到各孔', 'manual', '按比例设各孔 target→跳节点控制', '无+GET gates'],
        ['节点控制', '无', '跳转', '—'],
        ['调整仿真参数', '无', '跳虚拟仿真', '—'],
        ['采纳建议', '非 executing', '按 AI 推荐开度执行', '✓ POST /dispatch/execute'],
        ['忽略', '无', '关闭本次建议', '★ POST decisions/{id}/ignore'],
        ['L1/L2/L3', 'admin/算法工程师', '切自动级别', '★ PUT /dispatch/auto-level'],
    ])

    # ═══════════════════════════════════════
    h(doc, '5. 两套仿真对比（避免后端做混）')
    tbl(doc, ['', '虚拟仿真', '数字孪生', '节点控制'], [
        ['触发', '一键应用仿真', '开始仿真', '执行本孔/提交变更'],
        ['算数方', '前端公式', '后端引擎', '后端设备/PLC'],
        ['改真实设备', '否', '否（仿真域）', '是'],
        ['互锁', '无', '无', '有'],
        ['接口', '仅读 monitoring', '/v1/simulation/*', '★ gates/execute 待开发'],
    ])

    h(doc, '6. 开发优先级')
    tbl(doc, ['优先级', '项', '原因'], [
        ['P0', 'gates 字段对齐 + batch/单孔 execute', '节点控制核心，现全 Mock'],
        ['P0', 'execute 时互锁三规则', '与前端一致'],
        ['P1', 'mode / auto-level', '手动模式开关现 Mock'],
        ['P1', 'precheck 接口', '前后端互锁结果一致'],
        ['P2', '虚拟仿真 apply', '当前前端自给自足'],
        ['P2', 'simulation WS', '数字孪生体验'],
    ])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    main()

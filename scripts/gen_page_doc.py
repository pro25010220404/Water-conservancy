# -*- coding: utf-8 -*-
"""生成简洁版三页面操作说明 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime

DESKTOP = os.path.join(os.path.expanduser('~'), 'Desktop')
OUTPUT = os.path.join(DESKTOP, '三页面操作说明（简洁版）.docx')


def font(run, name='宋体', size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def cell_text(cell, text, bold=False):
    cell.text = ''
    font(cell.paragraphs[0].add_run(text), bold=bold)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell_text(t.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        font(r, '黑体', 14 if level == 1 else 12, True)


def p(doc, text):
    para = doc.add_paragraph()
    font(para.add_run(text))
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    font(para.add_run(text))


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(title.add_run('三页面操作说明（简洁版）'), '黑体', 18, True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(sub.add_run(f'向家坝闸门调度系统 · {datetime.now():%Y-%m-%d}'), size=11)

    h(doc, '0. 先看这个：三页各干什么')
    add_table(doc, ['页面', '菜单/路由', '干什么'], [
        ['虚拟仿真', '侧边栏 → 虚拟仿真\n/virtual-simulation', '调水位、调降雨，看流量/开度会怎么变'],
        ['节点控制', '调度决策 → 节点控制\n/dispatch/gates', '12 个孔逐个调开度，提交到设备（演示为模拟）'],
        ['数字孪生', '侧边栏 → 数字孪生\n/simulation', '3D/2D 看坝体，跑 AI 仿真、出报告'],
    ])
    p(doc, '推荐顺序：虚拟仿真调工况 → 节点控制微调开度 → 数字孪生看 3D 效果。')
    p(doc, '进节点控制前：先到「调度决策 → 运行控制」切成【手动】模式，否则滑块是灰的。')

    # ========== 虚拟仿真 ==========
    h(doc, '1. 虚拟仿真页')
    p(doc, '文件：VirtualSimulationPage.vue  路由：/virtual-simulation')

    h(doc, '1.1 页面上有什么', 2)
    bullet(doc, '左侧：水位滑块（110–200 m）、降雨滑块（0–120 mm/h）、操作按钮')
    bullet(doc, '右侧：上游/下游水位柱、入库/出库流量表、坝体 12 孔剖面图')
    bullet(doc, '顶部标签：「仿真生效中」= 已应用；「参数已锁定」= 滑块不可拖')

    h(doc, '1.2 每个按钮', 2)
    add_table(doc, ['按钮', '点了之后', '注意'], [
        ['常水位', '水位→175.7m，降雨→0', '锁定状态下点不了'],
        ['汛限偏高', '水位→185m，降雨→35', '同上'],
        ['枯水偏低', '水位→165m，降雨→0', '同上'],
        ['重置参数', '取消仿真，恢复接口实时数据', '其他页也会回到实时值'],
        ['锁定模拟', '锁住滑块，防误触', '再点一次变「解锁调节」'],
        ['一键应用仿真', '标记仿真生效，全站读仿真数据', '必须点这个，别的页才联动'],
        ['去节点控制调闸', '跳 /dispatch/gates', '没开手动模式会先跳运行控制'],
        ['节点控制（顶部小按钮）', '跳 /dispatch/gates', '不检查手动模式'],
        ['数字孪生（顶部小按钮）', '跳 /simulation', '—'],
        ['坝体上点某一孔', '选中该孔（仅本页高亮）', '不跳转'],
    ])
    p(doc, '拖滑块：未锁定时，拖动就会算一遍预览；但别的页面要等你点「一键应用仿真」才变。')

    h(doc, '1.3 数字怎么来的（简版）', 2)
    bullet(doc, '入库 ≈ 原入库 + 降雨×12 + (当前水位−原水位)×85')
    bullet(doc, '出库随水头（上下游水位差）变化')
    bullet(doc, '各孔开度按出库比例整体缩放')

    h(doc, '1.4 谁能进', 2)
    p(doc, 'operator / dispatcher / manager / admin / algorithm_engineer 均可。')

    # ========== 节点控制 ==========
    h(doc, '2. 节点控制页')
    p(doc, '文件：DispatchGatesPage.vue  路由：/dispatch/gates')

    h(doc, '2.1 进页面前', 2)
    bullet(doc, '运行控制页切【手动】— 否则所有调节按钮灰色')
    bullet(doc, '11#、12# 底孔默认【离线】— 能点选看数据，滑块/执行不可用')
    bullet(doc, '表孔 1#–8# 在线；9#、10# 中孔在线')

    h(doc, '2.2 顶部操作栏', 2)
    add_table(doc, ['按钮/开关', '作用', '细节'], [
        ['互锁严格 ↔ 手动绕过', '关=违规不能提交；开=可强制提交（弹确认）', '演示包默认「手动绕过」开'],
        ['虚拟仿真', '跳 /virtual-simulation', '—'],
        ['数字孪生', '跳 /simulation', '—'],
        ['表孔拉齐', '1#–8# 目标开度改成一样', '解决「对称性互锁」报错的最快办法'],
        ['全部复位', '所有孔 目标=当前，清空待提交', '不会动设备真实开度，只取消本地修改'],
        ['提交变更 (N)', '把 N 个待改孔一次性下发', '需手动模式；互锁阻断时要么拉齐要么开绕过'],
    ])

    h(doc, '2.3 告警条里的链接', 2)
    add_table(doc, ['链接', '去哪', '什么时候出现'], [
        ['去运行控制', '/dispatch/control', '没开手动模式'],
        ['详情', '弹窗列互锁项', '有阻断/警告'],
        ['调参', '/virtual-simulation', '虚拟仿真正在联动时（绿色提示）'],
    ])

    h(doc, '2.4 选孔与调开度', 2)
    bullet(doc, '选孔：点坝体剖面某一列，或点左下「闸门节点」列表')
    bullet(doc, '调开度：右下详情区 — 滑块 / 数字框 / 上调·下调（步长 1%）')
    bullet(doc, '目标≠当前 时，Tab 上会出现数字角标，底部出现「待提交变更」表')

    h(doc, '2.5 右下详情区按钮', 2)
    add_table(doc, ['按钮', '作用', '禁用条件'], [
        ['执行本孔', '只执行当前这一孔（current→target）', '离线 / 非手动 / 目标=当前'],
        ['重置', '当前孔 target 改回 current', '—'],
        ['同组同步', '表孔/中孔/底孔同组其他孔 target 跟当前孔对齐', '需先选中该组某一孔'],
    ])

    h(doc, '2.6 互锁规则（提交/执行时检查）', 2)
    add_table(doc, ['规则', '限值', '不过会怎样'], [
        ['表孔对称性', '1#–8# 目标开度最大差 ≤ 10%', '【阻断】不能提交 → 点「表孔拉齐」或开「手动绕过」'],
        ['单步变化', '单孔一次改动 ≤ 10%', '【警告】可提交但会二次确认'],
        ['生态流量', '预估出库 ≥ 500 m³/s', '【警告】可提交但会二次确认'],
    ])
    p(doc, '开「手动绕过」后：阻断项也会弹框问「仍要执行吗？」，点继续才行。')

    h(doc, '2.7 其他细节', 2)
    bullet(doc, '每 5 秒自动刷新数据；你有未提交的 target 时会保留，不被刷掉')
    bullet(doc, '离开节点控制 Tab 且还有未提交变更 → 弹窗问是否离开')
    bullet(doc, '虚拟仿真生效时，本页水位/流量/开度显示绿色仿真值，不是实时接口值')

    h(doc, '2.8 谁能进', 2)
    p(doc, 'dispatcher / manager / admin / algorithm_engineer（operator 进不了调度模块）。')

    # ========== 数字孪生 ==========
    h(doc, '3. 数字孪生页')
    p(doc, '文件：SimulationPage.vue  路由：/simulation')

    h(doc, '3.1 和虚拟仿真的区别', 2)
    add_table(doc, ['', '虚拟仿真', '数字孪生'], [
        ['干什么', '手调水位降雨，快速看 KPI', '3D 场景 + AI 模型训练 + 正式仿真任务'],
        ['有没有互锁', '无', '无（互锁在节点控制做）'],
        ['数据从哪来', '前端公式算', '后端仿真 API + WebSocket'],
    ])

    h(doc, '3.2 主要按钮（本页不跳别的路由）', 2)
    add_table(doc, ['位置', '按钮', '干什么'], [
        ['中栏顶部', '2D 剖面 / 3D 场景 / 全景 BIM', '切换视图；BIM 开全景弹窗'],
        ['中栏', '双击 3D 视口', '开全景控制弹窗'],
        ['右栏 Tab', '仿真视图', '调仿真参数、看闸门'],
        ['右栏 Tab', 'AI 模型导入训练', '导入模型 → 激活 → 训练'],
        ['右栏 Tab', '方案评估报告', '先跑完仿真 → 生成报告 → 下载 PDF'],
        ['右栏 Tab', '历史故障复盘', '选记录 → 导入仿真'],
        ['场景库', '新建场景 / 刷新', '管理仿真场景；有运行记录的场景不能删'],
        ['场景库', '编辑 / 删除', '改场景名或删空场景'],
        ['全景弹窗', '开始 / 暂停 / 重置', '控制仿真任务状态'],
        ['泄洪监测列表', '点 1#–5# 孔', '页内选中，3D 高亮'],
    ])
    p(doc, '虚拟仿真已应用时：本页水位、开度优先显示仿真叠加值。')

    h(doc, '3.3 谁能进', 2)
    p(doc, 'operator / dispatcher / manager / admin / algorithm_engineer 均可。')

    # ========== 运行控制 ==========
    h(doc, '4. 运行控制（节点控制前置，必看）')
    p(doc, '路由：/dispatch/control  Tab：调度决策 → 运行控制')

    add_table(doc, ['按钮', '去哪 / 做什么'], [
        ['手动 / 自动', '切模式；节点控制必须【手动】或 L1'],
        ['调整仿真参数', '→ /virtual-simulation'],
        ['节点控制', '→ /dispatch/gates'],
        ['分配到各孔', '总开度按比例分到各孔，然后 → /dispatch/gates'],
        ['节点级精确控制', '→ /dispatch/gates'],
        ['决策分析', '→ /dispatch/analysis'],
    ])

    # ========== 页间跳转一览 ==========
    h(doc, '5. 页间跳转一览（查表用）')
    add_table(doc, ['从哪', '按钮文字', '到哪'], [
        ['虚拟仿真', '节点控制 / 去节点控制调闸', '/dispatch/gates'],
        ['虚拟仿真', '数字孪生', '/simulation'],
        ['节点控制', '虚拟仿真', '/virtual-simulation'],
        ['节点控制', '数字孪生', '/simulation'],
        ['节点控制', '去运行控制', '/dispatch/control'],
        ['节点控制', '调参', '/virtual-simulation'],
        ['运行控制', '调整仿真参数', '/virtual-simulation'],
        ['运行控制', '节点控制', '/dispatch/gates'],
        ['侧边栏', '虚拟仿真', '/virtual-simulation'],
        ['侧边栏', '数字孪生', '/simulation'],
        ['侧边栏', '调度决策', '/dispatch/control（默认）'],
    ])

    h(doc, '6. 常见问题（3 条）')
    bullet(doc, '滑块灰的 → 运行控制切手动；若是 11#/12# → 离线改不了，换别的孔')
    bullet(doc, '提交报互锁 → 点「表孔拉齐」；演示时可开「手动绕过」')
    bullet(doc, '改了又变回去 → 5 秒刷新；「全部复位」只清本地 target，不清设备')

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    main()

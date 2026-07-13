# -*- coding: utf-8 -*-
"""生成后端对接说明 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

DESKTOP = os.path.join(os.path.expanduser('~'), 'Desktop')
OUTPUT = os.path.join(DESKTOP, '三页面-后端对接说明.docx')


def font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def cell(c, text, bold=False):
    c.text = ''
    font(c.paragraphs[0].add_run(str(text)), bold=bold)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell(t.rows[0].cells[i], h, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell(t.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        font(r, '黑体', 13 if level == 1 else 11, True)


def p(doc, text):
    para = doc.add_paragraph()
    font(para.add_run(text))
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    font(para.add_run(text))


def code_block(doc, text):
    para = doc.add_paragraph()
    font(para.add_run(text), 'Consolas', 9)


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.8)
        s.right_margin = Cm(2.8)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(t.add_run('三页面 — 后端对接说明'), '黑体', 18, True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(sub.add_run(f'前端仓库 Water-conservancy · {datetime.now():%Y-%m-%d}'), size=10)

    h(doc, '0. 给后端的一句话')
    p(doc, '前端有三个新页面：虚拟仿真、节点控制、数字孪生。虚拟仿真目前纯前端算数；节点控制和数字孪生需要你们提供 REST 接口。下面按「页面 → 调什么接口 → 字段格式 → 业务规则」写，可直接照着实现。')

    h(doc, '0.1 全局约定')
    bullet(doc, 'BaseURL：VITE_API_BASE_URL，默认 /api')
    bullet(doc, '鉴权：Header Authorization: Bearer {token}')
    bullet(doc, '统一响应：{ code: 0, msg, data, success, trace_id }，code≠0 为失败')
    bullet(doc, '字段命名：后端 snake_case，前端 L2 层做 camelCase 映射')
    bullet(doc, '默认水库 reservoir_id = 1')
    bullet(doc, 'Mock 降级：接口失败时前端走 mockStore，联调时务必返回 code=0')

    # ===== 页面1 =====
    h(doc, '1. 虚拟仿真页（/virtual-simulation）')
    p(doc, '【重点】这页目前没有专属后端接口，全是前端 Pinia + 本地公式。')

    h(doc, '1.1 它读了哪些已有接口', 2)
    table(doc, ['接口', '用途', '轮询'], [
        ['GET /api/monitoring/realtime?reservoir_id=1', '初始化 baseline 水位/流量', '进页调 1 次'],
        ['GET /api/monitoring/gates?reservoir_id=1', '坝体 12 孔数据', '进页 + 调参后 refresh'],
    ])

    h(doc, '1.2 前端自己算什么（后端不用管，除非以后要落库）', 2)
    bullet(doc, '入参：upstreamLevel（110–200m）、rainfall（0–120 mm/h）')
    bullet(doc, '出参：inflowRate、outflowRate、downstreamLevel、aggregateOpening、gateScale')
    bullet(doc, '公式见 src/utils/virtualSimulationEngine.ts')
    bullet(doc, 'active=true 后，fetchRealtimeKpi / fetchGates 返回值会被 overlay 覆盖显示')

    h(doc, '1.3 后端如果要接（可选，当前未实现）', 2)
    p(doc, '前端没有 POST 虚拟仿真接口。若后端要接管，建议：')
    table(doc, ['建议接口', '说明'], [
        ['POST /api/v1/hydrology/virtual-sim/apply', 'body: { upstream_level, rainfall } → 返回衍生 KPI'],
        ['DELETE /api/v1/hydrology/virtual-sim', '重置为实时监测值'],
    ])

    # ===== 页面2 =====
    h(doc, '2. 节点控制页（/dispatch/gates）')
    p(doc, '这是后端工作量最大的页：读闸门列表 + 执行开度变更 + 互锁校验。')

    h(doc, '2.1 已对接 / 在用的接口', 2)
    table(doc, ['方法', '路径', '频率', '说明'], [
        ['GET', '/api/monitoring/gates?reservoir_id=1', '每 5s', '12 孔列表，见 2.2 字段'],
        ['GET', '/api/monitoring/realtime?reservoir_id=1', '每 5s', '上下游水位、入库流量等'],
        ['GET', '/api/v1/dispatch/decisions?page=1&page_size=50&reservoir_id=1', '进页', '拼 DispatchStatus'],
        ['GET', '/api/v1/dispatch/decisions/{id}', '进页', '决策详情'],
        ['POST', '/api/v1/dispatch/execute', '运行控制页', '聚合总开度执行，见 2.4'],
        ['GET', '/api/v1/dispatch/gate-actions', '可选', '闸门动作历史'],
    ])

    h(doc, '2.2 GET /monitoring/gates 响应字段（单条）', 2)
    code_block(doc, '''{
  "id": 1,
  "name": "1#",
  "code": "GATE-001",
  "status": "online",        // online | offline | executing
  "opening": 45,             // 当前开度 0-100
  "target_opening": 45,      // 目标开度 0-100
  "mode": "AI-DQN",          // 控制模式字符串
  "flow_rate": 98,           // 单孔流量 m³/s
  "last_action_at": "14:32"
}''')
    p(doc, '闸门分组（前端 inferGateGroup，后端也可直接返回 group）：')
    table(doc, ['code 编号', '分组'], [
        ['GATE-001 ~ GATE-008', 'surface 表孔'],
        ['GATE-009 ~ GATE-010', 'mid 中孔'],
        ['GATE-011 ~ GATE-012', 'bottom 底孔'],
    ])
    p(doc, 'Mock 数据约定：11#、12# status=offline；演示环境表孔 opening/target 一致且在线。')

    h(doc, '2.3 GET /monitoring/realtime 响应字段', 2)
    code_block(doc, '''{
  "upstream_level": "378.52",
  "downstream_level": "269.18",
  "inflow_rate": "6350",
  "outflow_rate": "5820",
  "gate_opening": "34.0",    // 聚合开度
  "power_output": "4119",
  "timestamp": "2026-07-12T10:00:00Z"
}''')
    p(doc, '注意：字符串数字，前端 parseFloat。')

    h(doc, '2.4 【待开发】节点级执行接口 — 当前前端纯 Mock', 2)
    p(doc, 'DispatchGatesPage 点「执行本孔」「提交变更」走的是 mockExecuteGate / mockBatchExecute，没有调后端。需要你们实现：')
    table(doc, ['建议接口', '请求体', '期望行为'], [
        ['POST /api/v1/dispatch/gates/{equipment_id}/execute',
         '{ "target_opening": 50, "operate_note": "optional" }',
         '单孔执行；校验互锁+手动模式；返回 command_id'],
        ['POST /api/v1/dispatch/gates/batch-execute',
         '{ "items": [{ "equipment_id": 1, "target_opening": 50 }, ...] }',
         '批量执行；任一 block 级互锁失败则整批拒绝（或返回 partial）'],
        ['POST /api/v1/dispatch/gates/precheck',
         '{ "items": [{ "equipment_id": 1, "target_opening": 50 }] }',
         '仅预检不执行；返回 violations 数组，前端可改为调后端而非本地算'],
    ])
    p(doc, '执行成功后：opening 更新为 target_opening，status 短暂 executing 再变 online，写 gate-actions 日志。')

    h(doc, '2.5 【待开发】运行模式接口 — 当前前端 Mock', 2)
    table(doc, ['前端函数', '现状', '建议接口'], [
        ['putDispatchMode(manual|auto)', '接口 throw，走 mock', 'PUT /api/v1/dispatch/mode  body: { mode }'],
        ['putAutoLevel(1|2|3)', '接口 throw，走 mock', 'PUT /api/v1/dispatch/auto-level  body: { level }'],
        ['postCancelExecute()', '未对接', 'POST /api/v1/dispatch/cancel'],
        ['postIgnoreDecision()', '未对接', 'POST /api/v1/dispatch/decisions/{id}/ignore'],
    ])
    p(doc, 'canManualControl 判定：mode=manual 或 autoLevel=1。后端应在 execute 接口二次校验。')

    h(doc, '2.6 互锁规则 — 后端 execute 时必须校验', 2)
    p(doc, '前端本地预检在 src/utils/gateControl.ts → precheckGateChanges。后端应实现相同逻辑（或调 settings 里配置的规则）：')
    table(doc, ['rule_code', '规则', '阈值', '级别', '不满足时'], [
        ['symmetry_constraint', '表孔对称性', '1#-8# 在线孔 target 最大差 ≤ 10%', 'block', '拒绝执行，HTTP 400 + violations'],
        ['rate_exceeded', '单步变化率', '|target - opening| ≤ 10%', 'warn', '可执行但返回 warnings'],
        ['min_discharge_guarantee', '生态流量', '预估出库 ≥ 500 m³/s', 'warn', '可执行但返回 warnings'],
    ])
    p(doc, '预检响应建议格式（与前端 GatePrecheckResult 对齐）：')
    code_block(doc, '''{
  "passed": false,
  "violations": [{
    "rule_code": "symmetry_constraint",
    "rule_name": "对称性约束",
    "severity": "block",
    "affected_gate_ids": [1,2,3],
    "message": "表孔开度差 15% 超过 10% 限制"
  }]
}''')
    p(doc, '互锁规则配置已有接口（设置页用，execute 时应读取生效规则）：')
    bullet(doc, 'GET /api/v1/settings/gate-interlock/rules')
    bullet(doc, 'GET /api/v1/settings/gate-interlock/logs')

    h(doc, '2.7 前端轮询与 pending 逻辑（对接时注意）', 2)
    bullet(doc, '每 5s GET gates + realtime')
    bullet(doc, '用户改了 target 但未提交时，轮询应保留用户 target（前端已做；后端若也存 target，需区分「待下发」与「已生效」）')
    bullet(doc, 'interlockManualOverride 是纯前端演示开关，生产环境不应存在；后端永远 enforce 互锁')

    # ===== 页面3 =====
    h(doc, '3. 数字孪生页（/simulation）')
    p(doc, '走 /api/v1/simulation/*，接口较完整。虚拟仿真 active 时，本页 KPI 读 overlay 值。')

    h(doc, '3.1 仿真任务接口', 2)
    table(doc, ['方法', '路径', '说明'], [
        ['GET', '/api/v1/simulation/scenarios', '场景列表 page/page_size/keyword'],
        ['POST', '/api/v1/simulation/scenarios', '创建场景'],
        ['PUT', '/api/v1/simulation/scenarios/{id}', '更新场景'],
        ['DELETE', '/api/v1/simulation/scenarios/{id}', '删除；usage_count>0 应拒绝'],
        ['POST', '/api/v1/simulation/start', '启动仿真，见 3.2'],
        ['POST', '/api/v1/simulation/{id}/pause', '暂停'],
        ['POST', '/api/v1/simulation/{id}/resume', '继续'],
        ['POST', '/api/v1/simulation/{id}/reset', '重置'],
        ['PUT', '/api/v1/simulation/{id}/gate', '调闸门开度 body: { gate_opening: 0-100 }'],
        ['GET', '/api/v1/simulation/{id}/result?aggregation=raw', '仿真结果'],
        ['POST', '/api/v1/simulation/{id}/report', '生成报告 body: report_type/format/include_charts'],
    ])

    h(doc, '3.2 POST /simulation/start 请求体', 2)
    code_block(doc, '''{
  "scenario_id": 1,
  "model_id": 2,
  "reservoir_id": 1,
  "duration": 3600,
  "speed": 1,
  "params": {
    "initial_water_level": 175.7,
    "inflow_rate": 1920,
    "gate_opening": 45
  }
}''')
    p(doc, '响应需含 simulation_id、status、ws_endpoint（WebSocket 推送进度）。')

    h(doc, '3.3 故障复盘', 2)
    table(doc, ['方法', '路径', '说明'], [
        ['GET', '/api/v1/simulation/incidents?page&page_size&reservoir_id', '故障列表'],
        ['GET', '/api/v1/simulation/incidents/{id}', '故障详情'],
        ['POST', '/api/v1/simulation/import-incident', 'body: { incident_id }'],
    ])

    h(doc, '3.4 AI 模型（复用 settings 接口）', 2)
    bullet(doc, 'GET /api/settings/models — 模型列表')
    bullet(doc, 'POST /api/settings/models/upload — multipart 上传')
    bullet(doc, 'PUT /api/settings/models/{id}/activate — 激活')

    h(doc, '3.5 运行控制已有接口（节点控制前置）', 2)
    table(doc, ['方法', '路径', 'body'], [
        ['POST', '/api/v1/dispatch/execute', '{ reservoir_id, target_opening, decision_id? } → { command_id }'],
    ])
    p(doc, '这是「总开度」执行，不是逐孔。分配到各孔后仍跳节点控制页逐孔提交。')

    # ===== 优先级 =====
    h(doc, '4. 后端开发优先级（建议）')
    table(doc, ['优先级', '接口', '原因'], [
        ['P0', 'GET /monitoring/gates + /realtime 字段对齐', '节点控制每 5s 轮询，字段错就全页崩'],
        ['P0', 'POST gates/batch-execute + 互锁校验', '节点控制核心，现在全是 Mock'],
        ['P1', 'PUT dispatch/mode + auto-level', '手动模式开关，现在 Mock'],
        ['P1', 'POST gates/precheck', '前后端互锁结果一致'],
        ['P2', '虚拟仿真 apply 接口', '当前前端自给自足，可后置'],
        ['P2', 'simulation WS 进度推送', '数字孪生体验'],
    ])

    h(doc, '5. 前端源码索引（方便后端查）')
    table(doc, ['文件', '内容'], [
        ['src/api/monitoring.ts', 'gates / realtime 请求与 Mock'],
        ['src/api/dispatchPage.ts', '调度决策 API'],
        ['src/api/simulation.ts', '数字孪生 API'],
        ['src/stores/dispatch.ts', '节点控制状态、mockExecuteGate'],
        ['src/stores/virtualSimulation.ts', '虚拟仿真 overlay'],
        ['src/utils/gateControl.ts', '互锁预检算法'],
        ['src/utils/virtualSimulationEngine.ts', '虚拟仿真计算公式'],
        ['src/types/dispatch.ts', '调度类型定义'],
        ['src/types/gateControl.ts', '闸门/互锁类型'],
    ])

    h(doc, '6. 联调检查清单')
    bullet(doc, 'gates 返回 12 条，code 连续 GATE-001~012，status 合法')
    bullet(doc, 'realtime 数字字段可 parseFloat')
    bullet(doc, 'execute 后 gates 里 opening 更新，gate-actions 有记录')
    bullet(doc, '表孔 target 差 >10% 时 execute 返回 block violation')
    bullet(doc, 'offline 孔 execute 返回 400')
    bullet(doc, 'mode=auto 且 level=2/3 时 execute 返回 403（仅 L1/手动可操作）')
    bullet(doc, 'simulation/start 返回 simulation_id + ws_endpoint')

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    main()
